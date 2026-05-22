from datetime import date
from pathlib import Path

from lanes_ceo.contracts import Artifact, CriticReview, Job

WEEKLY_REPORT_SECTIONS = [
    "本周已完成工作",
    "科研进度复盘",
    "实验数据/仿真结果（含图表）",
    "工作亮点",
    "现存问题",
    "失败案例及原因分析",
    "下周详细工作计划",
    "需要团队及导师协助解决的问题",
]

EMPTY_PHRASE_FLAGS = [
    "有待加强", "进一步研究", "继续努力", "积极推进",
    "认真完成", "努力改进", "高度重视", "持续关注",
    "总体良好", "基本完成", "大致符合",
]

# SEU thesis format constants
_FONT_BODY_CN = "宋体"       # SimSun — body Chinese
_FONT_BODY_EN = "Times New Roman"  # body English
_FONT_HEADING_CN = "黑体"    # SimHei — heading Chinese
_FONT_HEADING_EN = "Times New Roman"
_BODY_SIZE = 12              # 小四, in pt
_HEADING1_SIZE = 16          # 三号
_HEADING2_SIZE = 14          # 四号
_HEADING3_SIZE = 12          # 小四 bold
_LINE_SPACING = 1.5
_PAGE_MARGIN_CM = 2.5


class WeeklyReportWorkflow:
    role_group = "weekly_report"
    actor_name = "weekly-report-actor"
    critic_name = "weekly-report-critic"

    def run_actor(self, job: Job) -> Artifact:
        message = job.input.get("message", "生成本周周报")
        today = date.today()
        week_label = f"{today.year}W{today.isocalendar()[1]}"

        from lanes_ceo.workflows.utils import get_artifact_dir, llm_chat

        prompt = _build_weekly_report_prompt(message)
        llm_response = llm_chat(
            "你是一名博士研究生，需按东南大学毕业论文格式撰写周报。"
            "输出严格分8个章节，每章标题用【】括起来，内容详实（100-200字）。",
            prompt,
        ) or ""

        artifact_dir = get_artifact_dir("weekly_report")
        ts = today.strftime("%H%M%S")
        docx_path = artifact_dir / f"weekly-report-{week_label}-{ts}.docx"

        summary_text = llm_response if llm_response else f"周报 {week_label}: {message}"
        _write_docx(docx_path, week_label, summary_text)

        return Artifact(
            artifact_id=f"artifact-{job.job_id}",
            job_id=job.job_id,
            artifact_type="weekly_report",
            summary=summary_text,
            artifact_paths=[str(docx_path)],
            sources=["llm-generated", "user-input"],
            risks=["LLM 内容需人工核实", "图表需手动补充"],
            user_confirmations=[
                "请核验所有实验数据和图表",
                "请确认下周计划是否可执行",
            ],
        )

    def run_critic(self, job: Job, artifact: Artifact) -> CriticReview:
        issues: list[str] = []
        score = 90

        if artifact.artifact_paths:
            path = Path(artifact.artifact_paths[0])
            if not path.exists():
                issues.append("生成的 docx 文件不存在")
                score -= 30

        if not llm_response_has_all_sections(artifact.summary):
            issues.append("部分周报板块缺失")

        empty_count = len(detect_empty_phrases(artifact.summary))
        if empty_count > 2:
            issues.append(f"检测到 {empty_count} 处空话/套话表述")
            score -= empty_count * 5

        approved = len(issues) == 0
        return CriticReview(
            review_id=f"review-{job.job_id}",
            job_id=job.job_id,
            review_result="approved" if approved else "returned",
            score=max(score, 0),
            issues=issues,
            approved=approved,
            return_to_actor=not approved,
            handoff_note="周报质量合格" if approved else "请修正后重新提交",
        )


def _build_weekly_report_prompt(message: str) -> str:
    sections_text = "\n".join(
        f"{i + 1}. {s}" for i, s in enumerate(WEEKLY_REPORT_SECTIONS)
    )
    return (
        f"请根据以下主题生成一份完整的周报，严格按以下8个章节撰写，每章内容充实（100-200字）：\n\n"
        f"{sections_text}\n\n"
        f"主题：{message}\n\n"
        f"格式要求：使用正式学术语言，图表需附题注（如：图1-1 xxx），公式需编号（如：(1-1)）。"
    )


def _write_docx(
    path: Path,
    week_label: str,
    content: str,
) -> None:
    """Generate a .docx following SEU thesis format: 宋体/Times New Roman, 小四, 1.5 line spacing."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── page setup (SEU: 2.5cm margins, A4) ──
    section = doc.sections[0]
    section.top_margin = Cm(_PAGE_MARGIN_CM)
    section.bottom_margin = Cm(_PAGE_MARGIN_CM)
    section.left_margin = Cm(_PAGE_MARGIN_CM)
    section.right_margin = Cm(_PAGE_MARGIN_CM)

    # ── default style ──
    style = doc.styles["Normal"]
    style.font.name = _FONT_BODY_EN
    style.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT_BODY_CN)
    style.font.size = Pt(_BODY_SIZE)
    pf = style.paragraph_format
    pf.line_spacing = _LINE_SPACING

    # ── title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    _add_run(title_para, f"周报 — {week_label}", _FONT_HEADING_CN, Pt(22), bold=True,
             color=RGBColor(0x1A, 0x1A, 0x2E))

    # ── subtitle ──
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(18)
    today_str = date.today().isoformat()
    _add_run(sub_para, f"生成日期：{today_str}", _FONT_BODY_CN, Pt(10),
             color=RGBColor(0x88, 0x88, 0x88))

    # ── body: parse LLM output into sections ──
    if content and not content.startswith("[LLM"):
        import re
        # Split by 【Section】 markers or by double-newline groups
        section_blocks = re.split(r"\n*(?:【|\[)([^】\]]+)(?:】|\])\s*", content)
        # The split alternates: [delimiter_prefix, section_title, content, section_title, content, ...]
        if len(section_blocks) >= 3 and section_blocks[0].strip() == "":
            # Content before first section marker
            preamble = section_blocks[0].strip()
            if preamble:
                doc.add_paragraph(preamble)
            idx = 1
            while idx + 1 < len(section_blocks):
                sec_title = section_blocks[idx]
                sec_content = section_blocks[idx + 1].strip()
                _add_section(doc, sec_title, sec_content)
                idx += 2
        else:
            # Fallback: split by double newlines
            paragraphs = content.split("\n\n")
            for block in paragraphs:
                stripped = block.strip()
                if not stripped:
                    continue
                lines = stripped.split("\n", 1)
                heading_text = lines[0].lstrip("0123456789. )（）#")
                body_text = lines[1] if len(lines) > 1 else ""
                _add_section(doc, heading_text, body_text)
    else:
        for i, section in enumerate(WEEKLY_REPORT_SECTIONS):
            _add_section(doc, f"{i + 1}. {section}", "（待 LLM 填充 — 请在 .env 中配置 LANES_CEO_LLM_API_KEY）")

    doc.save(str(path))


def _add_section(doc, heading_text: str, body_text: str) -> None:
    """Add a section heading + body paragraph with SEU formatting."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # Heading
    heading_para = doc.add_paragraph()
    heading_para.paragraph_format.space_before = Pt(12)
    heading_para.paragraph_format.space_after = Pt(6)
    heading_para.paragraph_format.line_spacing = _LINE_SPACING
    _add_run(heading_para, heading_text, _FONT_HEADING_CN, Pt(_HEADING2_SIZE), bold=True,
             color=RGBColor(0x1A, 0x1A, 0x2E))

    # Body
    if body_text:
        body_para = doc.add_paragraph()
        body_para.paragraph_format.first_line_indent = Pt(_BODY_SIZE * 2)  # 2-char indent
        body_para.paragraph_format.line_spacing = _LINE_SPACING
        _add_run(body_para, body_text, _FONT_BODY_CN, Pt(_BODY_SIZE))


def _add_run(paragraph, text: str, cn_font: str, size, bold: bool = False,
             color=None) -> None:
    """Add a run with proper Chinese+English font pair."""
    from docx.oxml.ns import qn

    run = paragraph.add_run(text)
    run.font.size = size
    run.bold = bold
    run.font.name = _FONT_BODY_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), _FONT_BODY_EN)
    rFonts.set(qn("w:hAnsi"), _FONT_BODY_EN)
    if color:
        run.font.color.rgb = color


def llm_response_has_all_sections(summary: str) -> bool:
    """Check that at least 5 of the 8 required sections appear in the summary."""
    required_keywords = [
        ["已完成", "完成工作", "本周工作"],
        ["进度", "复盘"],
        ["实验", "仿真", "数据"],
        ["亮点", "进展", "突破"],
        ["问题", "困难", "瓶颈"],
        ["失败", "教训", "不足"],
        ["下周", "计划"],
        ["协助", "支持", "资源"],
    ]
    found = 0
    for keywords in required_keywords:
        if any(kw in summary for kw in keywords):
            found += 1
    return found >= 5


def detect_empty_phrases(text: str) -> list[str]:
    return [p for p in EMPTY_PHRASE_FLAGS if p in text]
