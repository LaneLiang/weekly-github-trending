"""ChecklistGenerator — generate submission checklist as .docx (WPS-compatible)."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from lanes_ceo.workflows.manuscript_tracker.checkers.base import CheckResult
from lanes_ceo.workflows.manuscript_tracker.config import DEFAULT_CONFIG
from lanes_ceo.workflows.manuscript_tracker.profiles import JournalProfile

logger = logging.getLogger("lanes_ceo.manuscript_tracker.checklist")


class ChecklistGenerator:
    """Generate a formatted submission checklist DOCX file.

    Uses python-docx directly to build a WPS-compatible .docx file
    containing all check items and their results.
    """

    def generate(
        self,
        results: list[CheckResult],
        profile: JournalProfile,
        project: "LaTeXProject",  # noqa: F821
        output_dir: Path | None = None,
        filename: str | None = None,
    ) -> Path | None:
        """Generate the checklist DOCX file.

        Args:
            results: List of CheckResult objects from all checkers.
            profile: Journal profile for context.
            project: LaTeX project for context.
            output_dir: Output directory (defaults to project dir).
            filename: Output filename (defaults to config value).

        Returns:
            Path to the generated .docx file, or None on failure.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
        except ImportError:
            logger.warning("python-docx not installed; skipping checklist generation")
            return None

        output_dir = output_dir or project.project_dir
        output_path = output_dir / (filename or DEFAULT_CONFIG.checklist_filename)

        doc = Document()

        # ── Page setup ──
        section = doc.sections[0]
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # ── Styles ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # ── Title ──
        title = doc.add_heading(f"投稿合规检查清单 — {profile.journal_name}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Info table
        info_table = doc.add_table(rows=4, cols=2, style="Table Grid")
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        info_data = [
            ("期刊", profile.journal_name),
            ("文章类型", project.article_type or profile.default_type),
            ("主文件", project.main_tex.name),
            ("检查时间", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            self._set_cell_shading(info_table.cell(i, 0), "E8E8E8")

        doc.add_paragraph()

        # ── Summary section ──
        doc.add_heading("检查概要", level=2)
        total = len(results)
        passed = sum(1 for r in results if r.status == "pass")
        failed = sum(1 for r in results if r.status == "fail")
        warning = sum(1 for r in results if r.status == "warn")
        skipped = sum(1 for r in results if r.status == "skip")

        summary_text = (
            f"共 {total} 项检查: {passed} 项通过, {failed} 项失败, "
            f"{warning} 项警告, {skipped} 项跳过"
        )
        p = doc.add_paragraph(summary_text)
        if failed > 0:
            self._set_run_color(p.runs[0] if p.runs else p.add_run(summary_text), "FF0000")

        doc.add_paragraph()

        # ── Results by checker ──
        for result in results:
            status_label = {
                "pass": "通过",
                "fail": "失败",
                "warn": "警告",
                "skip": "跳过",
            }.get(result.status, result.status.upper())

            status_color = {
                "pass": "008000",
                "fail": "FF0000",
                "warn": "FF8C00",
                "skip": "808080",
            }.get(result.status, "000000")

            heading_text = f"[{status_label}] {result.checker_name}"
            h = doc.add_heading(heading_text, level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(
                    *self._hex_to_rgb(status_color)
                )

            if result.summary:
                doc.add_paragraph(f"概要: {result.summary}", style="List Bullet")

            # Items table
            if result.items:
                self._add_items_table(doc, result)

            # Fix suggestions
            if result.fix_suggestions:
                doc.add_paragraph("修复建议:", style="List Bullet")
                for suggestion in result.fix_suggestions:
                    doc.add_paragraph(suggestion[:300], style="List Bullet 2")

            doc.add_paragraph()  # spacer

        # ── Footer ──
        doc.add_paragraph("—" * 40)
        doc.add_paragraph(
            f"本清单由 LANEs_CEO 投稿合规检查引擎自动生成于 "
            f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )

        # Save
        try:
            doc.save(str(output_path))
            logger.info("Checklist DOCX saved to %s", output_path)
            return output_path
        except Exception as exc:
            logger.error("Failed to save checklist DOCX: %s", exc)
            return None

    def _add_items_table(self, doc, result: CheckResult) -> None:
        """Add a table of check items for a single checker result."""
        if not result.items:
            return

        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        table = doc.add_table(rows=1, cols=4, style="Table Grid")
        table.alignment = 1  # center

        # Header row
        hdr_cells = table.rows[0].cells
        headers = ["编号", "检查项", "状态", "详情 / 建议"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            self._set_cell_shading(hdr_cells[i], "4472C4")
            for p in hdr_cells[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for item in result.items:
            row = table.add_row()
            cells = row.cells

            cells[0].text = item.code
            cells[0].paragraphs[0].runs[0].font.size = Pt(8) if cells[0].paragraphs[0].runs else None

            cells[1].text = item.description
            for p in cells[1].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

            status_text = {
                "pass": "通过",
                "fail": "失败",
                "warn": "警告",
                "skip": "跳过",
            }.get(item.status, item.status)

            cells[2].text = status_text
            status_colors = {
                "pass": "008000",
                "fail": "FF0000",
                "warn": "FF8C00",
                "skip": "808080",
            }
            color = status_colors.get(item.status, "000000")
            for p in cells[2].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*self._hex_to_rgb(color))
                    run.font.bold = True
                    run.font.size = Pt(9)

            detail_text = item.detail or ""
            if item.fix_suggestion:
                detail_text += f"\n建议: {item.fix_suggestion}"
            cells[3].text = detail_text[:500]
            for p in cells[3].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

        # Set column widths
        widths = [Inches(1.2), Inches(2.5), Inches(0.6), Inches(2.5)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    @staticmethod
    def _set_cell_shading(cell, color_hex: str) -> None:
        """Set cell background shading."""
        from docx.oxml.ns import qn
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(
            qn("w:shd"),
            {
                qn("w:fill"): color_hex,
                qn("w:val"): "clear",
            },
        )
        shading_elm.append(shading)

    @staticmethod
    def _set_run_color(run, color_hex: str) -> None:
        """Set text color on a run."""
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*ChecklistGenerator._hex_to_rgb(color_hex))

    @staticmethod
    def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
        """Convert hex color string to RGB tuple."""
        hex_color = hex_color.lstrip("#")
        return (
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
